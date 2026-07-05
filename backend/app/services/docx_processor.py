import re
import random
from pathlib import Path
from docx import Document

SEP = "<<<<PARA_BREAK_@@@>>>>"


def resolve_paths(input_path, output_path):
    input_path = Path(input_path)

    if not input_path.exists():
        raise FileNotFoundError(f"Input file not found: {input_path}")

    output_path = Path(output_path)

    output_path.parent.mkdir(parents=True, exist_ok=True)

    return input_path, output_path


def process_options_block(block_text):
    """
    Shuffle options inside one {...} block.
    """

    inner = block_text.replace(SEP, "\n")

    lines = [
        line.strip()
        for line in re.split(r"\r?\n", inner)
        if line.strip()
    ]

    if not lines:
        return "{\n}"

    cleaned = []

    for line in lines:

        line = re.sub(r"^[\u2022\-\*\s]*", "", line)

        if line.startswith(("=", "#", "~")):
            line = line[1:].strip()

        cleaned.append(line)

    random.shuffle(cleaned)

    rebuilt = [
        "# " + option
        for option in cleaned
    ]

    return "{\n" + "\n".join(rebuilt) + "\n}"


def replace_blocks_in_joined(joined_text):

    pattern = re.compile(r"\{(.*?)\}", re.DOTALL)

    replacements = []

    def repl(match):

        original = match.group(0)

        new_block = process_options_block(match.group(1))

        replacements.append(
            (
                original,
                new_block
            )
        )

        return new_block

    new_text, count = pattern.subn(
        repl,
        joined_text
    )

    return new_text, replacements, count


def process_paragraphs_container(paragraphs):

    texts = [p.text for p in paragraphs]

    if not texts:
        return 0, []

    joined = SEP.join(texts)

    new_joined, replacements, count = replace_blocks_in_joined(joined)

    if count == 0:
        return 0, []

    new_texts = new_joined.split(SEP)

    if len(new_texts) == len(texts):

        for paragraph, new_text in zip(paragraphs, new_texts):
            paragraph.text = new_text

    else:

        paragraphs[0].text = new_joined.replace(
            SEP,
            "\n"
        )

        for paragraph in paragraphs[1:]:
            paragraph.text = ""

    return count, replacements


def shuffle_docx(input_path: str, output_path: str):
    """
    FastAPI entry point.

    Parameters
    ----------
    input_path : str
        Uploaded DOCX path

    output_path : str
        Output DOCX path

    Returns
    -------
    dict
    """

    input_path, output_path = resolve_paths(
        input_path,
        output_path
    )

    doc = Document(str(input_path))

    total_blocks = 0

    replacements = []

    # Top-level paragraphs

    count, reps = process_paragraphs_container(
        list(doc.paragraphs)
    )

    total_blocks += count

    replacements.extend(reps)

    # Tables

    for table in doc.tables:

        for row in table.rows:

            for cell in row.cells:

                count, reps = process_paragraphs_container(
                    list(cell.paragraphs)
                )

                total_blocks += count

                replacements.extend(reps)

    doc.save(str(output_path))

    return {
        "status": "success",
        "processed_blocks": total_blocks,
        "output_file": str(output_path),
        "sample_replacements": replacements[:5]
    }